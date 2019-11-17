# -*- coding:utf-8 -*-
import docx
import sys
from multiprocessing import Pool
import os
import urllib
import pymongo
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
import os

def get_lyrics_from_db():
    dbclient = pymongo.MongoClient("mongodb://localhost:27017/")
    db = dbclient["feitsui"]
    coll = db.lyric
    cursor = coll.find()
    lyrics = []
    for doc in cursor:
        lyrics.append(doc)
    return lyrics

def create_lyric_page(doc, lyric):
    """
    创建新的页
    """    
    p = doc.add_heading(u"", level=0)
    title_p = p.add_run(lyric['title'])
    title_p.font.name = u"华文楷体"
    title_p._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
    infos = lyric.get('info')
    if infos:
        if infos.get(u'歌手'):
            p = doc.add_paragraph(u"")
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            singer_p = p.add_run(infos.get(u'歌手'))
            singer_p.font.color.rgb = RGBColor(0, 32, 96)
            singer_p.font.name = "Microsoft YaHei"
            singer_p._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft Yahei')

    if lyric.get('desc'):
        p = doc.add_paragraph(u"")
        desc_p = p.add_run(lyric['desc'])
        desc_p.font.name = "Microsoft YaHei"
        desc_p._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft Yahei')
        #desc.style.font.color.rgb = RGBColor(0, 0, 0)

    if infos:
        for info in infos:
            p = doc.add_paragraph(u"", style="ListBullet")
            info_p = p.add_run(info + u": ")
            info_p.font.bold = True
            info_p.font.name = "Microsoft YaHei"
            info_p._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft Yahei')
            item = " ".join(infos[info])
            info_p = p.add_run(item)
            info_p.font.name = "Microsoft YaHei"
            info_p._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft Yahei')

    p = doc.add_paragraph(u"")
    line_p = p.add_run(u"粤拼（Lyrics）")
    line_p.font.name = u"华文楷体"
    line_p._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')
    line_p.bold = True
    line_p.font.color.rgb = RGBColor(0x3c, 0x76, 0x3d)
    line_p.font.size = Pt(12)

    lyric_text = lyric.get('lyric')
    lines =  lyric_text.split("\n")

    i = 0
    j = 0
    for line in lines:
        if line:
            if i % 2 == 0:
                l = doc.add_paragraph(u"")
                chinese = l.add_run(lines[j] + "\n")
                chinese.font.name = "Microsoft Yahei"
                chinese._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft Yahei')
                try:
                    pinyin = l.add_run(lines[j+1])
                    pinyin.font.name = "Cambria"
                except Exception as e:
                    print (e, lyric)
            i += 1
        else:
            doc.add_paragraph(u"")
        j += 1

    doc.add_paragraph(u"")
    lables = lyric.get('label')
    label = " ".join(lables)
    if lables:
        p = doc.add_paragraph(u'')
        label_p = p.add_run(u"标签（TAG）：")
        label_p.font.name = u"幼圆"
        label_p = p.add_run(label)
        label_p.font.name = u"幼圆"
        label_p._element.rPr.rFonts.set(qn('w:eastAsia'), u'幼圆')
        label_p.font.color.rgb = RGBColor(0x33, 0x7a, 0xb7)

def gen_doc(lyric):
    url = lyric['url']    
    html_name = url.split("/")[-1]
    filename, _ = os.path.splitext(html_name)
    filename = urllib.parse.unquote(filename)+ ".docx"
    filename = lyric["title"] + "-" + filename.replace('/', '-')
    filename = filename.replace(":","-").replace("/","-") # 有些歌名字叫做“3/8”或者"00:08:00"，不能用来作为文件名。
    doc = docx.Document()    
    if not os.path.exists("lyrics/"+filename):
        print ("Creating {0}".format(filename))
        create_lyric_page(doc, lyric)
        doc.save("./"+filename)
        doc

def create_doc():
        
    lyrics = get_lyrics_from_db()
    pool = Pool(10)
    for lyric in lyrics:
        pool.apply_async(gen_doc, args=(lyric,))
    pool.close()
    pool.join()    

    
if __name__ == "__main__":
    create_doc()
    print ("All FINISHED!")
